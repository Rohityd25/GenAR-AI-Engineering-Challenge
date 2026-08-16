"""
Report Assembly Layer
Assembles approved report sections into final documents in Markdown, HTML, and DOCX formats.
Follows the exact section sequence specified by ReportConfig and flags unapproved sections visibly.
"""

import logging
from pathlib import Path
from typing import Dict, Optional
from src.config import ReportConfig
from src.reviewer import ReviewStore

logger = logging.getLogger(__name__)


class ReportAssembler:
    """
    Concatenates approved sections and outputs final publication documents.
    """

    def __init__(self, report_config: ReportConfig, review_store: ReviewStore):
        self.config = report_config
        self.review_store = review_store

    def assemble_markdown(self) -> str:
        """Assembles the full report in standard GitHub-flavored Markdown."""
        lines = [
            f"# {self.config.title}",
            f"**Product Name:** {self.config.product_name}  ",
            f"**Regulatory Standard:** {self.config.regulatory_framework}  ",
            f"**Document Version:** {self.config.version}  ",
            "",
            "---",
            "",
        ]

        sorted_sections = sorted(self.config.sections, key=lambda s: s.display_order)
        for idx, sec in enumerate(sorted_sections, start=1):
            lines.append(f"## {idx}. {sec.name}")
            lines.append("")

            review_state = self.review_store.reviews.get(sec.section_id)
            if not review_state:
                lines.append("> [!WARNING]")
                lines.append(f"> **Section Incomplete:** No content generated for section `{sec.name}`.")
                lines.append("")
            elif review_state.status == "flagged":
                lines.append("> [!CAUTION]")
                lines.append(f"> **Section Flagged for Revision:** {review_state.review_comments or 'Pending author updates.'}")
                lines.append("")
                lines.append(review_state.content)
                lines.append("")
            elif review_state.status == "pending_review":
                lines.append("> [!NOTE]")
                lines.append("> **Draft Status:** This section is currently pending formal medical review approval.")
                lines.append("")
                lines.append(review_state.content)
                lines.append("")
            else:
                # Approved
                lines.append(review_state.content)
                lines.append("")

            lines.append("---")
            lines.append("")

        return "\n".join(lines)

    def assemble_html(self, md_content: str) -> str:
        """Converts assembled Markdown to a clean, styled HTML report."""
        import markdown
        try:
            html_body = markdown.markdown(md_content, extensions=["tables", "fenced_code", "nl2br"])
        except Exception:
            # Fallback simple conversion if markdown package has issues
            html_body = f"<pre>{md_content}</pre>"

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{self.config.title} - {self.config.product_name}</title>
<style>
  body {{
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
    line-height: 1.6;
    color: #1a202c;
    max-width: 960px;
    margin: 0 auto;
    padding: 30px 20px;
    background-color: #f8fafc;
  }}
  .report-container {{
    background: #ffffff;
    padding: 40px;
    border-radius: 8px;
    box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1), 0 2px 4px -1px rgba(0,0,0,0.06);
  }}
  h1 {{
    color: #1e3a8a;
    border-bottom: 2px solid #3b82f6;
    padding-bottom: 12px;
    font-size: 26px;
  }}
  h2 {{
    color: #1e40af;
    margin-top: 32px;
    padding-bottom: 6px;
    border-bottom: 1px solid #e2e8f0;
    font-size: 20px;
  }}
  h3 {{
    color: #334155;
    margin-top: 20px;
    font-size: 16px;
  }}
  table {{
    width: 100%;
    border-collapse: collapse;
    margin: 18px 0;
    font-size: 14px;
  }}
  th, td {{
    padding: 10px 12px;
    border: 1px solid #cbd5e1;
    text-align: left;
  }}
  th {{
    background-color: #f1f5f9;
    font-weight: 600;
    color: #334155;
  }}
  tr:nth-child(even) {{
    background-color: #f8fafc;
  }}
  blockquote {{
    border-left: 4px solid #3b82f6;
    background-color: #eff6ff;
    padding: 12px 16px;
    margin: 16px 0;
    border-radius: 0 4px 4px 0;
    font-size: 14px;
  }}
  code {{
    background-color: #f1f5f9;
    padding: 2px 6px;
    border-radius: 4px;
    font-family: Consolas, Monaco, monospace;
    font-size: 13px;
  }}
  hr {{
    border: 0;
    border-top: 1px solid #e2e8f0;
    margin: 24px 0;
  }}
</style>
</head>
<body>
<div class="report-container">
{html_body}
</div>
</body>
</html>
"""
        return html

    def assemble_docx(self, md_content: str, output_path: str) -> None:
        """Assembles a Microsoft Word DOCX document."""
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = docx.Document()
        
        # Title
        title = doc.add_heading(self.config.title, level=0)
        title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

        meta_p = doc.add_paragraph()
        meta_p.add_run(f"Product: {self.config.product_name} | Standard: {self.config.regulatory_framework} | Version: {self.config.version}").italic = True

        for line in md_content.split("\n"):
            line_str = line.strip()
            if not line_str or line_str.startswith("# ") or line_str == "---":
                continue
            elif line_str.startswith("## "):
                doc.add_heading(line_str[3:], level=1)
            elif line_str.startswith("### "):
                doc.add_heading(line_str[4:], level=2)
            elif line_str.startswith("#### "):
                doc.add_heading(line_str[5:], level=3)
            elif line_str.startswith("- "):
                doc.add_paragraph(line_str[2:], style="List Bullet")
            elif line_str.startswith("|") and not line_str.startswith("| :"):
                # Simplified table paragraph in docx
                doc.add_paragraph(line_str, style="Normal")
            elif line_str.startswith("> "):
                p = doc.add_paragraph(line_str[2:])
                p.runs[0].font.italic = True
            else:
                doc.add_paragraph(line_str)

        doc.save(output_path)
        logger.info(f"DOCX report saved to {output_path}")

    def export_all(self, output_dir: str = "output") -> Dict[str, str]:
        """Assembles and saves Markdown, HTML, and DOCX reports."""
        out_path = Path(output_dir)
        out_path.mkdir(parents=True, exist_ok=True)

        md_content = self.assemble_markdown()
        md_file = out_path / "report_output.md"
        with open(md_file, "w", encoding="utf-8") as f:
            f.write(md_content)

        html_content = self.assemble_html(md_content)
        html_file = out_path / "report_output.html"
        with open(html_file, "w", encoding="utf-8") as f:
            f.write(html_content)

        docx_file = out_path / "report_output.docx"
        try:
            self.assemble_docx(md_content, str(docx_file))
        except Exception as e:
            logger.warning(f"DOCX generation failed: {e}")

        # Also copy primary report to root per submission guide
        root_md = Path("report_output.md")
        with open(root_md, "w", encoding="utf-8") as f:
            f.write(md_content)

        return {
            "markdown": str(md_file),
            "html": str(html_file),
            "docx": str(docx_file),
            "root_markdown": str(root_md),
        }
